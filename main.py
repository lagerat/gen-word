import sys
from copy import deepcopy
from datetime import datetime

from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import Qt

from docx import Document
from openpyxl import load_workbook
from docx.enum.text import WD_COLOR_INDEX

class TableModel(QtCore.QAbstractTableModel):
    def __init__(self, data):
        super(TableModel, self).__init__()
        self._data = data
        

    headerNames = ['Институт', 'Направление', 'Профиль', 'Семестр', 'Вид практики', 'Тип практики', 'Трудоемкость',
                   'Дата начала', 'Дата окончания', 'Компетенции']

    def data(self, index, role):
        if role == Qt.DisplayRole or role == Qt.EditRole:
            value = self._data[index.row()][index.column()]

            if isinstance(value, datetime):
                # Render time to YYY-MM-DD.
                return value.strftime("%Y-%m-%d")

            if isinstance(value, float):
                # Render float to 2 dp
                return "%.2f" % value

            if isinstance(value, str):
                # Render strings with quotes
                return '%s' % value

                # Default (anything not captured above: e.g. int)
            return value
        if role == Qt.ToolTipRole:
            if index.column() == 7 or index.column() == 8:
                return "Формат должен быть: ДД месяц ГГГГ"
    def rowCount(self, index):
        # The length of the outer list.
        return len(self._data)

    def columnCount(self, index):
        # The following takes the first sub-list, and returns
        # the length (only works if all rows are an equal length)
        return len(self._data[0])

    def headerData(self, section, orientation, role):
        # section is the index of the column/row.
        if role == Qt.DisplayRole:
            if orientation == Qt.Horizontal:
                return self.headerNames[section]

    def setData(self, index, value, role):
        if role == Qt.EditRole:
            if index.column() == 7 or index.column() == 8:
                self._data[index.row()][index.column()] = value
                return True
            return False
        return False

    def flags(self, index):
        return Qt.ItemIsSelectable | Qt.ItemIsEnabled | Qt.ItemIsEditable


class ValidatedItemDelegate(QtWidgets.QItemDelegate):
    def createEditor(self, widget, option, index):
        if not index.isValid():
            return 0
        if index.column() == 7 or index.column() == 8:
            editor = QtWidgets.QLineEdit(widget)
            validator = QtGui.QRegExpValidator(QtCore.QRegExp("\d{2} [a-zA-Zа-яА-Я]{0,10} \d{4}"), editor)
            editor.setValidator(validator)
            return editor

        return super(ValidatedItemDelegate, self).createEditor(widget, option, index)

class Ui_MainWindow(object):
    data = [
        ['', '', '', '', '', '', '', '', '', '']
    ]

    def setupUi(self, MainWindow):
        MainWindow.setObjectName("Word Generator")
        MainWindow.resize(1280, 900)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout.setObjectName("verticalLayout")
        self.fisrtStageTable = QtWidgets.QTableView(self.centralwidget)
        self.fisrtStageTable.setObjectName("fisrtStageTable")

        self.tableModel = TableModel(self.data)
        self.fisrtStageTable.setModel(self.tableModel)

        self.verticalLayout.addWidget(self.fisrtStageTable)

        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setObjectName("horizontalLayout")

        self.uploadBtn = QtWidgets.QPushButton(self.centralwidget)
        self.uploadBtn.setObjectName("uploadBtn")
        self.horizontalLayout.addWidget(self.uploadBtn)

        self.createBtn = QtWidgets.QPushButton(self.centralwidget)
        self.createBtn.setObjectName("createBtn")
        self.horizontalLayout.addWidget(self.createBtn)
        self.verticalLayout.addLayout(self.horizontalLayout)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 800, 22))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)
        self.fisrtStageTable.setItemDelegate(ValidatedItemDelegate())
        self.connectFunctions()

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "Word Generator"))
        self.uploadBtn.setText(_translate("MainWindow", "Загрузить"))
        self.createBtn.setText(_translate("MainWindow", "Сгененировать"))

    def connectFunctions(self):
        self.uploadBtn.clicked.connect(self.onUploadBtn_clicked)
        self.createBtn.clicked.connect(self.onCreateBtn_clicked)

    def __get_rows_xy(self, plan):
        pred = lambda row: row[0].value and (row[0].value.startswith("Блок 2") or row[0].value.startswith("Блок 3"))
        block1, block2 = filter(pred, plan.iter_rows())

        return (block1[0].row + 2, block2[0].row - 1)

    def __proceed_table(self, document_name):
        document = load_workbook(document_name, read_only = True)

        title = document["Титул"]
        plan = document["План"]
        competencies = document["Компетенции"]
 
        first_three_cells = [title["D38"].value, title["D27"].value, title["D30"].value]

        x, y = self.__get_rows_xy(plan)

        for row in plan[f"D{x}:F{y}"]:
            for cell in filter(lambda c: c.value, row):
                for ch in cell.value:
                    curr_row = first_three_cells.copy()

                    curr_row.append(ch)

                    if plan["B" + str(cell.row)].value[-2] == "У":
                        curr_row.append("Учебная")
                    else:
                        curr_row.append("Производственная")

                    curr_row.append(plan["C" + str(cell.row)].value)
                    curr_row.append(plan["N" + str(cell.row)].value + "/" + plan["L" + str(cell.row)].value)

                    comps_l = [s.strip() for s in plan["CP" + str(cell.row)].value.split(";")]
                    comps_str = str()

                    for row in filter(lambda r: r[1].value and r[1].value in comps_l, competencies.iter_rows()):
                        comps_str += " " + row[1].value + " - " + row[3].value
                        if comps_str[-1] != ";":
                            comps_str += ";"

                    curr_row.extend(["", ""])                    

                    curr_row.append(comps_str.lstrip())

                    self.__addRecordToTable(curr_row)

    def onUploadBtn_clicked(self):
        self.uploadBtn.setDisabled(True)

        filePaths = QtWidgets.QFileDialog.getOpenFileNames(None,"Выберите файлы","","Excel File (*.xlsx *.xls)")

        fileNames = filePaths[0]

        for name in fileNames:
            self.__proceed_table(name)

        self.uploadBtn.setEnabled(True)   

    def __fill_run(self, run, row):
        s = run.text.strip()

        if not s.isdigit():
            return

        i = int(s)

        if i < 1 or i > 10:
            return

        run.font.highlight_color = WD_COLOR_INDEX.AUTO

        if i == 4:
            sem = int(row[i - 1])
            run.text = str(sem // 2 + sem % 2)
        else:
            run.text = row[i - 1]

    def __fill_doc(self, doc, row):
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.highlight_color == WD_COLOR_INDEX.RED:
                    self.__fill_run(run, row)

        for table in doc.tables:
            for r in table.rows:
                for cell in r.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.highlight_color == WD_COLOR_INDEX.RED:
                                self.__fill_run(run, row)
                    

    def onCreateBtn_clicked(self):

        docsPath = QtWidgets.QFileDialog.getOpenFileNames(None, "Выберите файлы в качестве примера", "", "Word (*.docx *.docm *.doc )")

        docsNames = docsPath[0]

        if len(docsNames) == 0:
            return        

        originalDocs = [Document(name) for name in docsNames]

        for  rowData in self.data:
            if rowData[7] != '' and rowData[8] != '':
                docs = [deepcopy(doc) for doc in originalDocs]

                for doc in docs:
                    self.__fill_doc(doc, rowData)

                for idx, doc in enumerate(docs):
                    doc.save(str(idx) + ".docx")

    def __addRecordToTable(self, record):
        lenghtList = len(record)
        if lenghtList < 10:
            return 0
        if self.data[0][0] == '':
            self.data.clear()
        self.data.append(record)
        self.fisrtStageTable.model().layoutChanged.emit()
        self.fisrtStageTable.resizeColumnsToContents()

if __name__ == "__main__":
    import sys

    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()

    ui.setupUi(MainWindow)

    MainWindow.show()

    sys.exit(app.exec())
