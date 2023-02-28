import os
import sys
from copy import deepcopy
from datetime import datetime
from datetime import date

from operator import itemgetter

from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import Qt


from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from openpyxl.utils import get_column_letter
from openpyxl import load_workbook, Workbook
from openpyxl.worksheet.table import Table, TableStyleInfo
from delegate import ValidatedItemDelegate
from delegate import DateEditDelegate
from tableModel import TableModel


class Ui_MainWindow(object):
    data = []
    selectedCells = []
    codeDirection = {
        '01.03.02': 'Прикладная математика и информатика',
        '02.03.02': 'Фундаментальная информатика и информационные технологии',
        '09.03.01': 'Информатика и вычислительная техника',
        '09.03.03': 'Прикладная информатика',
        '10.03.01': 'Информационная безопасность',
        '38.03.05': 'Бизнес-информатика',
        '42.03.01': 'Реклама и связи с общественностью',
        '09.03.02': 'Информационные системы и технологии',
        '11.03.01': 'Радиотехника',
        '11.03.02': 'Инфокоммуникационные технологии и системы связи',
        '11.03.03': 'Конструирование и технология электронных средств',
        '11.03.04': 'Электроника и наноэлектроника',
        '20.03.01': 'Техносферная безопасность',
        '10.05.02': 'Информационная безопасность телекоммуникационных систем',
        '11.05.01': 'Радиоэлектронные системы и комплексы',
        '11.05.02': 'Специальные радиотехнические системы',
        '11.05.04': 'Инфокоммуникационные технологии и системы специальной связи'
    }

    def setupUi(self, MainWindow):
        MainWindow.setObjectName("Word Generator")
        MainWindow.resize(1280, 900)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")

        self.verticalLayout_2 = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout_2.setObjectName("verticalLayout_2")

        self.horizontalLayout_2 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_2.setObjectName("horizontalLayout_2")

        self.fisrtStageTable = QtWidgets.QTableView(self.centralwidget)
        self.fisrtStageTable.setObjectName("fisrtStageTable")

        self.horizontalLayout_2.addWidget(self.fisrtStageTable)

        self.verticalLayout = QtWidgets.QVBoxLayout()
        self.verticalLayout.setObjectName("verticalLayout")

        self.clearCellBtn = QtWidgets.QPushButton(self.centralwidget)
        self.clearCellBtn.setObjectName("clearCellBtn")

        self.verticalLayout.addWidget(self.clearCellBtn)

        self.deleteRowBtn = QtWidgets.QPushButton(self.centralwidget)
        self.deleteRowBtn.setObjectName("deleteRowBtn")

        self.verticalLayout.addWidget(self.deleteRowBtn)

        self.stretchBtn = QtWidgets.QPushButton(self.centralwidget)
        self.stretchBtn.setEnabled(False)
        self.stretchBtn.setObjectName("stretchBtn")

        self.verticalLayout.addWidget(self.stretchBtn)

        self.horizontalLayout_2.addLayout(self.verticalLayout)
        self.verticalLayout_2.addLayout(self.horizontalLayout_2)

        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setObjectName("horizontalLayout")

        self.uploadBtn = QtWidgets.QPushButton(self.centralwidget)
        self.uploadBtn.setObjectName("uploadBtn")

        self.horizontalLayout.addWidget(self.uploadBtn)

        self.createBtn = QtWidgets.QPushButton(self.centralwidget)
        self.createBtn.setObjectName("createBtn")

        self.horizontalLayout.addWidget(self.createBtn)
        self.verticalLayout_2.addLayout(self.horizontalLayout)

        MainWindow.setCentralWidget(self.centralwidget)

        self.menubar = QtWidgets.QMenuBar(MainWindow)

        self.menubar.setGeometry(QtCore.QRect(0, 0, 800, 22))
        self.menubar.setObjectName("menubar")
        self.menu = QtWidgets.QMenu(self.menubar)
        self.menu.setObjectName("menu")
        MainWindow.setMenuBar(self.menubar)

        self.uploadToExcel = QtWidgets.QAction(MainWindow)
        self.uploadToExcel.setObjectName("uploadToExcel")
        self.infoMenu = QtWidgets.QAction(MainWindow)
        self.infoMenu.setObjectName("infoMenu")

        self.menu.addAction(self.uploadToExcel)
        self.menu.addAction(self.infoMenu)
        self.menubar.addAction(self.menu.menuAction())

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

        self.tableModel = TableModel(self.data)
        self.fisrtStageTable.setModel(self.tableModel)
        self.fisrtStageTable.setVerticalScrollMode(QtWidgets.QAbstractItemView.ScrollPerPixel)
        self.fisrtStageTable.setHorizontalScrollMode(QtWidgets.QAbstractItemView.ScrollPerPixel)

        self.fisrtStageTable.setItemDelegateForColumn(7, DateEditDelegate(MainWindow))
        self.fisrtStageTable.setItemDelegateForColumn(8, DateEditDelegate(MainWindow))
        self.connectFunctions()

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "Word Generator"))
        self.clearCellBtn.setText(_translate("MainWindow", "Очистить"))
        self.deleteRowBtn.setText(_translate("MainWindow", "Удалить строку"))
        self.stretchBtn.setText(_translate("MainWindow", "Расстянуть"))
        self.uploadBtn.setText(_translate("MainWindow", "Загрузить"))
        self.createBtn.setText(_translate("MainWindow", "Сгененировать"))
        self.menu.setTitle(_translate("MainWindow", "Меню"))
        self.uploadToExcel.setText(_translate("MainWindow", "Выгрузить в excel"))
        self.infoMenu.setText(_translate("MainWindow", "Коды для word файлов"))

    def connectFunctions(self):
        self.uploadBtn.clicked.connect(self.onUploadBtn_clicked)
        self.createBtn.clicked.connect(self.onCreateBtn_clicked)
        self.uploadToExcel.triggered.connect(self.uploadExcelFunc)
        self.infoMenu.triggered.connect(self.onInfoMenuClick)
        self.fisrtStageTable.selectionModel().selectionChanged.connect(self.tableSelectionChangend)
        self.clearCellBtn.clicked.connect(self.onClearCellBtn_click)
        self.deleteRowBtn.clicked.connect(self.onDelRowBtn_click)
        self.stretchBtn.clicked.connect(self.onStretchBtn_click)

    def onInfoMenuClick(self):
        msg = QtWidgets.QMessageBox()
        msg.setIcon(QtWidgets.QMessageBox.Information)
        msg.setText("1 - Институт\n"
                    "2 - Направление\n"
                    "3 - Профиль\n"
                    "4 - Семестр\n"
                    "5 - Вид практики\n"
                    "6 - Тип Практики\n"
                    "7 - Трудоёмкость\n"
                    "8 - Дата начала\n"
                    "9 - Дата окончания\n"
                    "10 - Компетенции\n"
                    "11 - Текущий год\n"
                    "12 - Сокращенине института\n"
                    "18 - Последний будний день от даты начала\n"
                    "19 - Последний будний день от даты окончания")
        msg.setWindowTitle("Коды")
        msg.exec()

    def onStretchBtn_click(self):
        if len(self.selectedCells) > 0:
            firstRecord = self.data[self.selectedCells[0][0]][self.selectedCells[0][1]]
            for record in self.selectedCells[1:]:
                self.data[record[0]][record[1]] = firstRecord

    def onDelRowBtn_click(self):
        if len(self.selectedCells) > 0:
            self.selectedCells = sorted(self.selectedCells, key=itemgetter(0), reverse=True)

            for record in self.selectedCells:
                del self.data[record[0]]
            self.fisrtStageTable.model().layoutChanged.emit()
            self.fisrtStageTable.resizeColumnsToContents()
            self.selectedCells.clear()

    def onClearCellBtn_click(self):
        if len(self.selectedCells) > 0:
            for record in self.selectedCells:
                self.data[record[0]][record[1]] = ''
            self.fisrtStageTable.model().layoutChanged.emit()
            self.fisrtStageTable.resizeColumnsToContents()

    def tableSelectionChangend(self, selected, deselected):
        for ix in selected.indexes():
            self.selectedCells.append([ix.row(), ix.column()])

        for ix in deselected.indexes():
            count = 0
            while count < len(self.selectedCells):
                if (self.selectedCells[count][0] == ix.row() and self.selectedCells[count][1] == ix.column()):
                    del self.selectedCells[count]
                else:
                    count += 1
        if len(self.selectedCells) > 1:
            isSretchable = True
            length = len(self.selectedCells)
            for i in range(1, length):
                if self.selectedCells[i][1] != self.selectedCells[i - 1][1]:
                    isSretchable = False
                    break
            if isSretchable:
                self.stretchBtn.setEnabled(True)
            else:
                self.stretchBtn.setEnabled(False)
        else:
            self.stretchBtn.setEnabled(False)

    def uploadExcelFunc(self):
        doc = Workbook()

        active = doc.active

        isEmpty = True

        active.append(self.tableModel.headerNames)

        maxLengths = [len(name) for name in self.tableModel.headerNames]

        for row in self.data:
            if row[7].toString() != "" and row[8].toString() != "":
                temp_row = row[:7] + [row[7].toString("dd.MM.yyyy"), row[8].toString("dd.MM.yyyy"), row[9]]

                active.append(temp_row)

                isEmpty = False

                for i, maxLength in enumerate(maxLengths):
                    maxLengths[i] = max(maxLength, len(temp_row[i]))

        if isEmpty:
            return

        for i, maxLength in enumerate(maxLengths):
            active.column_dimensions[get_column_letter(i + 1)].width = maxLength + 5

        try:
            doc.save("Сводная таблица.xlsx")
        except:
            None

    def __get_rows_xy(self, plan):
        pred = lambda row: row[0].value and (row[0].value.startswith("Блок 2") or row[0].value.startswith("Блок 3"))
        block1, block2 = filter(pred, plan.iter_rows())

        return (block1[0].row + 2, block2[0].row - 1)

    def __proceed_table(self, document_name):
        document = load_workbook(document_name, read_only=True)

        title = document["Титул"]
        plan = document["План"]
        competencies = document["Компетенции"]

        first_three_cells = [title["D38"].value, title["D27"].value, title["D30"].value]
        direction = self.codeDirection.get(first_three_cells[1])
        if direction != None:
            first_three_cells[1] = first_three_cells[1] + "-" + direction

        x, y = self.__get_rows_xy(plan)

        for row in plan[f"D{x}:F{y}"]:
            for cell in filter(lambda c: c.value, row):
                for ch in cell.value:
                    curr_row = first_three_cells.copy()

                    curr_row.append(ch)

                    if plan["B" + str(cell.row)].value[-2] == "У":
                        curr_row.append("Учебная")
                    else:
                        curr_row.append("Производственная")

                    curr_row.append(plan["C" + str(cell.row)].value)
                    curr_row.append(plan["N" + str(cell.row)].value + "/" + plan["L" + str(cell.row)].value)

                    comps_l = [s.strip() for s in plan["CP" + str(cell.row)].value.split(";")]
                    comps_str = str()

                    for row in filter(lambda r: r[1].value and r[1].value in comps_l, competencies.iter_rows()):
                        comps_str += " " + row[1].value + " - " + row[3].value
                        if comps_str[-1] != ";":
                            comps_str += ";"

                    curr_row.extend([QtCore.QDate(), QtCore.QDate()])

                    curr_row.append(comps_str.lstrip())

                    self.__addRecordToTable(curr_row)

    def onUploadBtn_clicked(self):
        self.uploadBtn.setDisabled(True)

        filePaths = QtWidgets.QFileDialog.getOpenFileNames(None, "Выберите файлы", "", "Excel File (*.xlsx *.xls)")

        fileNames = filePaths[0]

        for name in fileNames:
            self.__proceed_table(name)

        self.uploadBtn.setEnabled(True)

    def __fill_run(self, run, row):
        s = run.text.strip()

        if not s.isdigit():
            return

        i = int(s)

        if (i < 1 or i > 12) and i != 18 and i != 19:
            return

        run.font.highlight_color = WD_COLOR_INDEX.AUTO

        if i == 4:
            sem = int(row[i - 1])
            run.text = str(sem // 2 + sem % 2)
        elif i == 8 or i == 9:
            day = row[i - 1].toString("\"dd\"").lower()
            month = row[i - 1].toString("MMMM").lower()
            year = row[i - 1].toString("yyyy г.").lower()
            if month[-1] == 'т':
                finalDate = day + " " + month + 'a' + " " + year
            else:
                finalDate = day + " " + month[:-1] + 'я' + " " + year
            run.text = finalDate
        elif i == 11:
            currentYear = date.today().year
            run.text = str(currentYear)
        elif i == 12:
            run.text = str("".join(word[0].upper() for word in row[0].split()))
        elif i == 18 or i == 19:
            newDate = QtCore.QDate.currentDate()
            newDate.setDate(row[i - 11].year(), row[i - 11].month(), row[i - 11].day())
            while (QtCore.QDate.dayOfWeek(newDate) > 5):
                newDate = newDate.addDays(-1)
            day = newDate.toString("\"dd\"").lower()
            month = newDate.toString("MMMM").lower()
            year = newDate.toString("yyyy г.").lower()
            if month[-1] == 'т':
                finalDate = day + " " + month + 'a' + " " + year
            else:
                finalDate = day + " " + month[:-1] + 'я' + " " + year
            run.text = finalDate
        else:
            run.text = row[i - 1]

    def __fill_Comptencies(self, run, row, table):
        s = run.text.strip()

        if not s.isdigit():
            return

        i = int(s)

        if i < 1 or i > 12:
            return
        run.font.highlight_color = WD_COLOR_INDEX.AUTO
        competencies = str(row[i - 1]).split(';')
        run.text = competencies[0]
        length = len(competencies)
        if length > 1:
            for x in range(1, length):
                if competencies[x] == '':
                    continue
                newrow = table.add_row().cells
                newrow[0].text = competencies[x]

    def __fill_doc(self, doc, row):
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.highlight_color == WD_COLOR_INDEX.RED:
                    self.__fill_run(run, row)

        for table in doc.tables:
            for r in table.rows:
                for cell in r.cells:
                    for paragraph in cell.paragraphs:
                        s = ''
                        for run in paragraph.runs:
                            s = s + run.text
                            if not s.isdigit():
                                s = ''
                            if run.font.highlight_color == WD_COLOR_INDEX.RED:
                                if s == "10":
                                    self.__fill_Comptencies(run, row, table)
                                else:
                                    self.__fill_run(run, row)

    def onCreateBtn_clicked(self):
        docsPath = QtWidgets.QFileDialog.getOpenFileNames(None, "Выберите файлы в качестве примера", "",
                                                          "Word (*.docx *.docm *.doc )")

        docsNames = docsPath[0]

        if len(docsNames) == 0:
            return

        originalDocs = [Document(name) for name in docsNames]

        absWorkingDir = os.path.abspath(os.getcwd())

        for rowData in self.data:
            if rowData[7].toString() != '' and rowData[8].toString() != '':
                docs = [deepcopy(doc) for doc in originalDocs]

                for doc in docs:
                    self.__fill_doc(doc, rowData)

                for idx, doc in enumerate(docs):
                    onlyDirection = rowData[1][:8]

                    path = os.path.join(absWorkingDir, onlyDirection, rowData[2], rowData[3])

                    os.makedirs(path, exist_ok=True)

                    doc.save(os.path.join(path, "_".join(
                        [onlyDirection, rowData[2], rowData[3], os.path.basename(docsNames[idx])])))

    def __addRecordToTable(self, record):
        self.data.append(record)
        self.fisrtStageTable.model().layoutChanged.emit()
        self.fisrtStageTable.resizeColumnsToContents()
